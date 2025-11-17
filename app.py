import streamlit as st
import requests
import matplotlib.pyplot as plt
import pandas as pd 
from collections import Counter
from fpdf import FPDF
from docx import Document
import io
import numpy as np
from sklearn.cluster import KMeans
from sklearn.preprocessing import StandardScaler
import seaborn as sns
from datetime import datetime, timedelta
import re

# ---------- Inject Custom CSS ----------
st.markdown("""
<style>
/* light blue Sidebar Gradient */
[data-testid="stSidebar"] {
    background: linear-gradient(135deg, #A0D8E8, #3B8EA5);
    padding: 1.5rem;
    min-width: 330px;
    max-width: 340px;
    color: white;
    box-shadow: 2px 0 10px rgba(0, 0, 0, 0.1);
    overflow-x: hidden;
}

/* Force all text in sidebar to be white */
[data-testid="stSidebar"] * {
    color: white !important;
}

/* Sidebar Header Section */
[data-testid="stSidebar"] > div:first-child {
    padding: 1.5rem 1rem;
    border-radius: 12px;
    text-align: center;
}

/* Navigation title */
[data-testid="stSidebar"] h1, 
[data-testid="stSidebar"] h2 {
    font-size: 22px;
    font-weight: bold;
}

/* Section label */
[data-testid="stSidebar"] label {
    font-size: 17px;
    margin-bottom: 10px;
}

/* Radio buttons */
.css-1n76uvr div[data-baseweb="radio"] label {
    background-color: rgba(255, 255, 255, 0.1);
    border-radius: 6px;
    padding: 10px 14px;
    margin-bottom: 6px;
    cursor: pointer;
    transition: all 0.3s ease;
    white-space: nowrap;
    overflow: hidden;
    text-overflow: ellipsis;
    display: block;
}

/* Hover and Active Radio */
.css-1n76uvr div[data-baseweb="radio"] label:hover {
    background-color: rgba(255, 255, 255, 0.25);
    transform: translateX(4px);
}
.css-1n76uvr div[data-baseweb="radio"] input:checked + div {
    background-color: white !important;
    color: #9F44D3 !important;
    font-weight: bold;
}

/* Light Theme for Main App Area */
body, .main, .block-container {
    background-color: #f9f9f9;
    color: #333333;
}
h1, h2, h3, h4 {
    color: #222222;
}

/* AI Feature Cards */
.ai-card {
    background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
    padding: 20px;
    border-radius: 10px;
    color: white;
    margin: 10px 0;
}
</style>
""", unsafe_allow_html=True)

# ---------- GitHub API Helpers ----------
def fetch_profile(username):
    url = f"https://api.github.com/users/{username}"
    r = requests.get(url)
    return r.json() if r.status_code == 200 else None

def fetch_repos(username):
    url = f"https://api.github.com/users/{username}/repos?per_page=100"
    r = requests.get(url)
    if r.status_code != 200:
        return [], {}, {}, {}
    repos = r.json()
    langs, stars, links = [], {}, {}
    for repo in repos:
        if repo["language"]:
            langs.append(repo["language"])
        stars[repo["name"]] = repo["stargazers_count"]
        links[repo["name"]] = repo["html_url"]
    return repos, Counter(langs), stars, links

def fetch_commits(username, repo_name):
    """Fetch recent commits for activity analysis"""
    url = f"https://api.github.com/repos/{username}/{repo_name}/commits?per_page=100"
    r = requests.get(url)
    return r.json() if r.status_code == 200 else []

# -------- NEW AI/ML FEATURES --------

def ml_developer_clustering(profiles_data):
    """
    ML Feature 1: Cluster developers based on their GitHub metrics
    Uses K-Means clustering to group similar developers
    """
    if len(profiles_data) < 3:
        return None, None
    
    features = []
    usernames = []
    
    for username, data in profiles_data.items():
        profile = data['profile']
        features.append([
            profile.get('followers', 0),
            profile.get('following', 0),
            profile.get('public_repos', 0),
            len(data.get('languages', {})),
            sum(data.get('stars', {}).values())
        ])
        usernames.append(username)
    
    # Normalize features
    scaler = StandardScaler()
    features_scaled = scaler.fit_transform(features)
    
    # K-Means clustering
    n_clusters = min(3, len(profiles_data))
    kmeans = KMeans(n_clusters=n_clusters, random_state=42)
    clusters = kmeans.fit_predict(features_scaled)
    
    return clusters, usernames

def ai_skill_prediction(lang_counter, repos, profile):
    """
    AI Feature 2: Predict developer expertise level and specialization
    Uses rule-based AI with scoring system
    """
    score = 0
    skills = []
    expertise_level = "Beginner"
    
    # Language diversity score
    lang_count = len(lang_counter)
    score += min(lang_count * 10, 50)
    
    # Repository quality score
    repo_count = profile.get('public_repos', 0)
    score += min(repo_count * 2, 40)
    
    # Star rating (popularity score)
    total_stars = sum([repo.get('stargazers_count', 0) for repo in repos])
    if total_stars > 100:
        score += 30
    elif total_stars > 50:
        score += 20
    elif total_stars > 10:
        score += 10
    
    # Follower influence score
    followers = profile.get('followers', 0)
    if followers > 100:
        score += 20
    elif followers > 50:
        score += 10
    
    # Determine expertise level
    if score >= 80:
        expertise_level = "Expert"
    elif score >= 50:
        expertise_level = "Advanced"
    elif score >= 25:
        expertise_level = "Intermediate"
    
    # Skill specialization detection
    if "Python" in lang_counter and lang_counter["Python"] > 5:
        skills.append("Python Specialist")
    if "JavaScript" in lang_counter or "TypeScript" in lang_counter:
        skills.append("Web Developer")
    if "Java" in lang_counter or "Kotlin" in lang_counter:
        skills.append("Enterprise Developer")
    if "C++" in lang_counter or "C" in lang_counter or "Rust" in lang_counter:
        skills.append("Systems Programmer")
    
    # Check for ML/AI indicators
    ml_keywords = ['machine-learning', 'deep-learning', 'neural', 'tensorflow', 'pytorch', 'ai']
    for repo in repos:
        desc = (repo.get('description') or '').lower()
        topics = repo.get('topics', [])
        if any(kw in desc or kw in ' '.join(topics) for kw in ml_keywords):
            if "AI/ML Engineer" not in skills:
                skills.append("AI/ML Engineer")
            break
    
    return {
        'expertise_level': expertise_level,
        'score': score,
        'specializations': skills if skills else ["General Developer"],
        'recommendations': generate_recommendations(score, lang_counter, repos)
    }

def generate_recommendations(score, lang_counter, repos):
    """Generate personalized recommendations for improvement"""
    recommendations = []
    
    if score < 50:
        recommendations.append("🎯 Create more public repositories to showcase your work")
        recommendations.append("📚 Contribute to open-source projects to gain visibility")
    
    if len(lang_counter) < 3:
        recommendations.append("💡 Learn additional programming languages to diversify skills")
    
    total_stars = sum([repo.get('stargazers_count', 0) for repo in repos])
    if total_stars < 10:
        recommendations.append("⭐ Improve documentation and README files to attract more stars")
    
    if not any(repo.get('description') for repo in repos[:5]):
        recommendations.append("📝 Add detailed descriptions to your repositories")
    
    return recommendations

def ai_contribution_pattern_analysis(username, repos):
    """
    AI Feature 3: Analyze contribution patterns and predict activity trends
    """
    activity_score = 0
    patterns = {
        'consistency': 'Unknown',
        'activity_level': 'Unknown',
        'prediction': 'Unknown'
    }
    
    # Analyze recent updates
    recent_repos = [r for r in repos if r.get('updated_at')]
    if recent_repos:
        dates = [datetime.strptime(r['updated_at'], '%Y-%m-%dT%H:%M:%SZ') for r in recent_repos]
        latest_activity = max(dates)
        days_since_activity = (datetime.now() - latest_activity).days
        
        if days_since_activity < 7:
            patterns['activity_level'] = 'Very Active'
            patterns['consistency'] = 'Consistent'
            activity_score = 100
        elif days_since_activity < 30:
            patterns['activity_level'] = 'Active'
            patterns['consistency'] = 'Regular'
            activity_score = 75
        elif days_since_activity < 90:
            patterns['activity_level'] = 'Moderate'
            patterns['consistency'] = 'Occasional'
            activity_score = 50
        else:
            patterns['activity_level'] = 'Inactive'
            patterns['consistency'] = 'Rare'
            activity_score = 25
        
        # Predict future activity
        if activity_score >= 75:
            patterns['prediction'] = 'Expected to remain highly active'
        elif activity_score >= 50:
            patterns['prediction'] = 'Likely to continue moderate activity'
        else:
            patterns['prediction'] = 'May become more active or remain dormant'
    
    return patterns, activity_score

def ml_repository_recommender(user_langs, all_repos):
    """
    ML Feature 4: Recommend repositories based on language similarity
    Simple collaborative filtering approach
    """
    recommendations = []
    
    for repo in all_repos:
        repo_lang = repo.get('language')
        if repo_lang in user_langs:
            score = user_langs[repo_lang] * repo.get('stargazers_count', 0)
            recommendations.append({
                'name': repo['name'],
                'url': repo['html_url'],
                'stars': repo['stargazers_count'],
                'score': score,
                'language': repo_lang
            })
    
    # Sort by score and return top 5
    recommendations.sort(key=lambda x: x['score'], reverse=True)
    return recommendations[:5]

def sentiment_analysis_simple(text):
    """
    Simple sentiment analysis for repository descriptions
    """
    positive_words = ['awesome', 'great', 'excellent', 'powerful', 'efficient', 'innovative', 
                      'modern', 'simple', 'fast', 'robust', 'reliable', 'scalable']
    negative_words = ['deprecated', 'old', 'slow', 'complex', 'buggy', 'legacy']
    
    text_lower = text.lower()
    pos_count = sum(1 for word in positive_words if word in text_lower)
    neg_count = sum(1 for word in negative_words if word in text_lower)
    
    if pos_count > neg_count:
        return 'Positive', pos_count - neg_count
    elif neg_count > pos_count:
        return 'Negative', neg_count - pos_count
    else:
        return 'Neutral', 0

def calculate_developer_score(profile, repos, lang_counter):
    """
    AI Feature 5: Calculate comprehensive developer score using multiple metrics
    """
    metrics = {
        'code_quality': 0,
        'community_impact': 0,
        'consistency': 0,
        'diversity': 0
    }
    
    # Code Quality (based on stars and forks)
    total_stars = sum([r.get('stargazers_count', 0) for r in repos])
    total_forks = sum([r.get('forks_count', 0) for r in repos])
    metrics['code_quality'] = min((total_stars * 0.7 + total_forks * 0.3) / 10, 100)
    
    # Community Impact
    followers = profile.get('followers', 0)
    metrics['community_impact'] = min(followers * 2, 100)
    
    # Consistency (repos updated recently)
    recent_count = sum(1 for r in repos if (datetime.now() - 
                      datetime.strptime(r['updated_at'], '%Y-%m-%dT%H:%M:%SZ')).days < 90)
    metrics['consistency'] = min((recent_count / max(len(repos), 1)) * 100, 100)
    
    # Diversity (language variety)
    metrics['diversity'] = min(len(lang_counter) * 15, 100)
    
    # Overall score
    overall = sum(metrics.values()) / 4
    
    return metrics, overall

# -------- REPORT GENERATORS --------
def generate_pdf_report(username, profile, repos):
    def safe_text(text):
        return text.encode("latin-1", "ignore").decode("latin-1")
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=safe_text(f"GitHub Report for {username}"), ln=True, align='C')
    pdf.ln(10)
    for key in ["name", "location", "public_repos", "followers", "following"]:
        value = profile.get(key, "N/A")
        pdf.cell(200, 10, txt=safe_text(f"{key.capitalize()}: {value}"), ln=True)
    pdf.ln(5)
    pdf.cell(200, 10, txt="Repositories:", ln=True)
    for repo in repos[:10]:
        pdf.cell(200, 10, txt=safe_text(f"- {repo['name']}: Stars {repo['stargazers_count']}"), ln=True)
    path = f"{username}_report.pdf"
    pdf.output(path)
    return path

def generate_docx_report(username, profile, repos):
    doc = Document()
    doc.add_heading(f"GitHub Report for {username}", 0)
    for key in ["name", "location", "public_repos", "followers", "following"]:
        doc.add_paragraph(f"{key.capitalize()}: {profile.get(key, 'N/A')}")
    doc.add_heading("Repositories", level=1)
    for repo in repos[:10]:
        doc.add_paragraph(f"- {repo['name']}: ⭐ {repo['stargazers_count']}")
    path = f"{username}_report.docx"
    doc.save(path)
    return path

def generate_resume(profile, repos):
    doc = Document()
    doc.add_heading(f"{profile.get('name', profile['login'])}'s Developer Resume", 0)
    doc.add_paragraph(f"GitHub: {profile['html_url']}")
    doc.add_paragraph(f"Location: {profile.get('location', 'N/A')}")
    doc.add_paragraph(f"Public Repos: {profile.get('public_repos')}, Followers: {profile.get('followers')}")
    doc.add_heading("Top Projects", level=1)
    for repo in repos[:5]:
        doc.add_paragraph(f"{repo['name']} - {repo['description'] or 'No description'}")
    path_docx = f"{profile['login']}_resume.docx"
    doc.save(path_docx)

    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt=f"{profile.get('name', profile['login'])}'s Developer Resume", ln=True, align='C')
    pdf.ln(10)
    pdf.cell(200, 10, txt=f"GitHub: {profile['html_url']}", ln=True)
    pdf.cell(200, 10, txt=f"Location: {profile.get('location', 'N/A')}", ln=True)
    pdf.cell(200, 10, txt=f"Public Repos: {profile.get('public_repos')}, Followers: {profile.get('followers')}", ln=True)
    pdf.ln(5)
    pdf.cell(200, 10, txt="Top Projects:", ln=True)
    for repo in repos[:5]:
        pdf.cell(200, 10, txt=f"- {repo['name']}: {repo['description'] or 'No description'}", ln=True)
    path_pdf = f"{profile['login']}_resume.pdf"
    pdf.output(path_pdf)
    return path_docx, path_pdf

# ---------- Session State ----------
if 'visualize_data' not in st.session_state:
    st.session_state.visualize_data = None
if 'ai_analysis_cache' not in st.session_state:
    st.session_state.ai_analysis_cache = {}

# ---------- Streamlit UI ----------
st.set_page_config(page_title="GitHub Profile Analyzer", layout="wide")
st.title("🐙 GitHub Profile Analyzer with AI/ML")

st.sidebar.header("🔍 Navigation")
action = st.sidebar.radio("Choose a section:", [
    "Analyze One User", "🤖 AI Skill Assessment", "📊 ML Developer Clustering",
    "🎯 Smart Recommendations", "📈 Activity Prediction", "Visualization", 
    "Compare Two Users", "Trending Repositories", "Trending Developers", "Resume Generator"
])

# ---------- Analyze One User ----------
if action == "Analyze One User":
    username = st.text_input("Enter GitHub Username")
    if username:
        profile = fetch_profile(username)
        repos, lang_counter, stars, links = fetch_repos(username)

        if profile:
            st.session_state.visualize_data = (profile, lang_counter)
            st.session_state.ai_analysis_cache[username] = {
                'profile': profile,
                'repos': repos,
                'languages': lang_counter,
                'stars': stars
            }

            col1, col2 = st.columns([1, 3])
            with col1:
                st.image(profile["avatar_url"], width=150)
            with col2:
                st.subheader(f"👤 {profile.get('name') or profile['login']}")
                st.markdown(f"📍 {profile.get('location', 'N/A')}")
                st.markdown(f"📦 Public Repos: {profile['public_repos']}")
                st.markdown(f"👥 Followers: {profile['followers']}, Following: {profile['following']}")
                st.markdown(f"[🔗 GitHub Profile]({profile['html_url']})")

            # Quick AI Insights
            st.markdown("---")
            st.subheader("🤖 Quick AI Insights")
            col1, col2, col3 = st.columns(3)
            
            metrics, overall_score = calculate_developer_score(profile, repos, lang_counter)
            with col1:
                st.metric("Overall Score", f"{overall_score:.1f}/100")
            with col2:
                skill_data = ai_skill_prediction(lang_counter, repos, profile)
                st.metric("Expertise Level", skill_data['expertise_level'])
            with col3:
                patterns, activity_score = ai_contribution_pattern_analysis(username, repos)
                st.metric("Activity Level", patterns['activity_level'])

            if lang_counter:
                st.subheader("🧠 Languages Used")
                col1, col2 = st.columns(2)
                with col1:
                    fig1, ax1 = plt.subplots()
                    ax1.pie(lang_counter.values(), labels=lang_counter.keys(), autopct='%1.1f%%')
                    ax1.axis('equal')
                    st.pyplot(fig1)
                with col2:
                    fig2, ax2 = plt.subplots()
                    ax2.bar(lang_counter.keys(), lang_counter.values(), color='skyblue')
                    ax2.set_ylabel("Count")
                    ax2.set_title("Languages Frequency")
                    st.pyplot(fig2)

            st.subheader("📊 Top Starred Repos")
            for name, star in sorted(stars.items(), key=lambda x: x[1], reverse=True)[:5]:
                st.markdown(f"- [{name}]({links[name]}): ⭐ {star}")

            st.subheader("📥 Download Report")
            df = pd.DataFrame(repos)
            st.download_button("Download CSV", df.to_csv(index=False), file_name=f"{username}_report.csv")
            pdf_path = generate_pdf_report(username, profile, repos)
            with open(pdf_path, "rb") as f:
                st.download_button("Download PDF", f, file_name=pdf_path)
            doc_path = generate_docx_report(username, profile, repos)
            with open(doc_path, "rb") as f:
                st.download_button("Download DOCX", f, file_name=doc_path)

# ---------- AI Skill Assessment ----------
elif action == "🤖 AI Skill Assessment":
    st.subheader("🤖 AI-Powered Skill Assessment")
    username = st.text_input("Enter GitHub Username for AI Analysis")
    
    if username:
        if username in st.session_state.ai_analysis_cache:
            data = st.session_state.ai_analysis_cache[username]
        else:
            profile = fetch_profile(username)
            repos, lang_counter, stars, _ = fetch_repos(username)
            data = {'profile': profile, 'repos': repos, 'languages': lang_counter, 'stars': stars}
            st.session_state.ai_analysis_cache[username] = data
        
        if data['profile']:
            skill_assessment = ai_skill_prediction(data['languages'], data['repos'], data['profile'])
            
            st.markdown(f"""
            <div class="ai-card">
                <h2>🎓 Expertise Level: {skill_assessment['expertise_level']}</h2>
                <h3>📊 AI Score: {skill_assessment['score']}/100</h3>
            </div>
            """, unsafe_allow_html=True)
            
            st.subheader("🎯 Specializations Detected")
            for spec in skill_assessment['specializations']:
                st.success(f"✅ {spec}")
            
            st.subheader("💡 AI-Generated Recommendations")
            for rec in skill_assessment['recommendations']:
                st.info(rec)
            
            # Detailed metrics
            st.subheader("📈 Detailed Performance Metrics")
            metrics, overall = calculate_developer_score(data['profile'], data['repos'], data['languages'])
            
            col1, col2 = st.columns(2)
            with col1:
                st.metric("Code Quality", f"{metrics['code_quality']:.1f}/100")
                st.metric("Community Impact", f"{metrics['community_impact']:.1f}/100")
            with col2:
                st.metric("Consistency", f"{metrics['consistency']:.1f}/100")
                st.metric("Diversity", f"{metrics['diversity']:.1f}/100")
            
            # Radar chart
            fig, ax = plt.subplots(figsize=(8, 6), subplot_kw=dict(projection='polar'))
            categories = list(metrics.keys())
            values = list(metrics.values())
            values += values[:1]
            angles = np.linspace(0, 2 * np.pi, len(categories), endpoint=False).tolist()
            angles += angles[:1]
            
            ax.plot(angles, values, 'o-', linewidth=2, color='#667eea')
            ax.fill(angles, values, alpha=0.25, color='#764ba2')
            ax.set_xticks(angles[:-1])
            ax.set_xticklabels(categories)
            ax.set_ylim(0, 100)
            ax.set_title("Developer Performance Radar", pad=20)
            st.pyplot(fig)

# ---------- ML Developer Clustering ----------
elif action == "📊 ML Developer Clustering":
    st.subheader("📊 ML-Based Developer Clustering")
    st.write("Enter multiple GitHub usernames to cluster similar developers using Machine Learning")
    
    usernames_input = st.text_area("Enter usernames (one per line)", height=150)
    
    if st.button("🔍 Cluster Developers"):
        usernames = [u.strip() for u in usernames_input.split('\n') if u.strip()]
        
        if len(usernames) >= 3:
            profiles_data = {}
            progress_bar = st.progress(0)
            
            for idx, username in enumerate(usernames):
                profile = fetch_profile(username)
                if profile:
                    repos, lang_counter, stars, _ = fetch_repos(username)
                    profiles_data[username] = {
                        'profile': profile,
                        'repos': repos,
                        'languages': lang_counter,
                        'stars': stars
                    }
                progress_bar.progress((idx + 1) / len(usernames))
            
            if len(profiles_data) >= 3:
                clusters, usernames_list = ml_developer_clustering(profiles_data)
                
                st.success(f"✅ Successfully clustered {len(profiles_data)} developers!")
                
                # Display clusters
                cluster_df = pd.DataFrame({
                    'Username': usernames_list,
                    'Cluster': [f"Group {c+1}" for c in clusters]
                })
                
                st.subheader("🎯 Clustering Results")
                st.dataframe(cluster_df)
                
                # Visualization
                fig, ax = plt.subplots(figsize=(10, 6))
                for i in range(max(clusters) + 1):
                    cluster_users = [u for u, c in zip(usernames_list, clusters) if c == i]
                    y_pos = list(range(len(cluster_users)))
                    ax.barh(y_pos, [i+1]*len(cluster_users), label=f'Group {i+1}')
                    for j, user in enumerate(cluster_users):
                        ax.text(i+0.5, j, user, ha='center', va='center')
                
                ax.set_xlabel('Cluster Group')
                ax.set_title('Developer Clustering Visualization')
                ax.legend()
                st.pyplot(fig)
            else:
                st.error("Not enough valid profiles fetched. Please check usernames.")
        else:
            st.warning("Please enter at least 3 usernames for clustering.")

# ---------- Smart Recommendations ----------
elif action == "🎯 Smart Recommendations":
    st.subheader("🎯 AI-Powered Smart Recommendations")
    username = st.text_input("Enter GitHub Username")
    
    if username:
        if username in st.session_state.ai_analysis_cache:
            data = st.session_state.ai_analysis_cache[username]
        else:
            profile = fetch_profile(username)
            repos, lang_counter, stars, _ = fetch_repos(username)
            data = {'profile': profile, 'repos': repos, 'languages': lang_counter, 'stars': stars}
        
        if data['profile']:
            st.subheader("📚 Recommended Repositories")
            recommendations = ml_repository_recommender(data['languages'], data['repos'])
            
            for rec in recommendations:
                st.markdown(f"""
                **[{rec['name']}]({rec['url']})**
                - Language: {rec['language']} | Stars: ⭐ {rec['stars']}
                - Relevance Score: {rec['score']:.2f}
                """)
            
            st.subheader("📊 Repository Sentiment Analysis")
            sentiments = []
            for repo in data['repos'][:10]:
                desc = repo.get('description', '')
                if desc:
                    sentiment, strength = sentiment_analysis_simple(desc)
                    sentiments.append({
                        'Repository': repo['name'],
                        'Sentiment': sentiment,
                        'Strength': strength
                    })
            
            if sentiments:
                sent_df = pd.DataFrame(sentiments)
                st.dataframe(sent_df)

# ---------- Activity Prediction ----------
elif action == "📈 Activity Prediction":
    st.subheader("📈 Activity Pattern Analysis & Prediction")
    username = st.text_input("Enter GitHub Username")
    
    if username:
        if username in st.session_state.ai_analysis_cache:
            data = st.session_state.ai_analysis_cache[username]
        else:
            profile = fetch_profile(username)
            repos, lang_counter, stars, _ = fetch_repos(username)
            data = {'profile': profile, 'repos': repos, 'languages': lang_counter}
        
        if data['profile']:
            patterns, activity_score = ai_contribution_pattern_analysis(username, data['repos'])
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Activity Level", patterns['activity_level'])
            with col2:
                st.metric("Consistency", patterns['consistency'])
            with col3:
                st.metric("Activity Score", f"{activity_score}/100")
            
            st.info(f"🔮 **Prediction:** {patterns['prediction']}")
            
            # Activity timeline
            if data['repos']:
                dates = []
                for repo in data['repos']:
                    if repo.get('updated_at'):
                        try:
                            date = datetime.strptime(repo['updated_at'], '%Y-%m-%dT%H:%M:%SZ')
                            dates.append(date)
                        except:
                            pass
                
                if dates:
                    df_dates = pd.DataFrame({'Date': dates})
                    df_dates['Month'] = df_dates['Date'].dt.to_period('M')
                    activity_by_month = df_dates.groupby('Month').size()
                    
                    st.subheader("📅 Activity Timeline")
                    fig, ax = plt.subplots(figsize=(12, 5))
                    activity_by_month.plot(kind='bar', ax=ax, color='#667eea')
                    ax.set_xlabel('Month')
                    ax.set_ylabel('Repository Updates')
                    ax.set_title('Monthly Activity Pattern')
                    plt.xticks(rotation=45)
                    st.pyplot(fig)

# ---------- Visualization ----------
elif action == "Visualization":
    if st.session_state.visualize_data:
        profile, lang_counter = st.session_state.visualize_data
        st.subheader("📊 Profile Overview")
        st.markdown(f"👤 **{profile.get('name') or profile['login']}**")

        data = {
            "Followers": profile['followers'],
            "Following": profile['following'],
            "Public Repos": profile['public_repos']
        }
        st.bar_chart(pd.DataFrame.from_dict(data, orient='index', columns=["Count"]))

        if lang_counter:
            st.subheader("🧠 Languages Used")
            col1, col2 = st.columns(2)
            with col1:
                fig1, ax1 = plt.subplots()
                ax1.pie(lang_counter.values(), labels=lang_counter.keys(), autopct='%1.1f%%')
                ax1.axis('equal')
                st.pyplot(fig1)
            with col2:
                fig2, ax2 = plt.subplots()
                ax2.bar(lang_counter.keys(), lang_counter.values(), color='orange')
                ax2.set_ylabel("Count")
                ax2.set_title("Languages Frequency")
                st.pyplot(fig2)
    else:
        st.info("Go to 'Analyze One User' first to load a profile.")

# ---------- Compare Two Users ----------
elif action == "Compare Two Users":
    col1, col2 = st.columns(2)
    with col1:
        user1 = st.text_input("GitHub Username 1")
    with col2:
        user2 = st.text_input("GitHub Username 2")
    if user1 and user2:
        profile1 = fetch_profile(user1)
        profile2 = fetch_profile(user2)
        repos1, langs1, stars1, _ = fetch_repos(user1)
        repos2, langs2, stars2, _ = fetch_repos(user2)

        if profile1 and profile2:
            # AI Comparison
            st.subheader("🤖 AI Skill Comparison")
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown(f"### {user1}")
                skill1 = ai_skill_prediction(langs1, repos1, profile1)
                st.metric("Expertise", skill1['expertise_level'])
                st.metric("AI Score", f"{skill1['score']}/100")
                st.write("**Specializations:**")
                for spec in skill1['specializations']:
                    st.write(f"- {spec}")
            
            with col2:
                st.markdown(f"### {user2}")
                skill2 = ai_skill_prediction(langs2, repos2, profile2)
                st.metric("Expertise", skill2['expertise_level'])
                st.metric("AI Score", f"{skill2['score']}/100")
                st.write("**Specializations:**")
                for spec in skill2['specializations']:
                    st.write(f"- {spec}")

            st.subheader("📈 Total Stars Comparison")
            st.bar_chart(pd.DataFrame({user1: [sum(stars1.values())], user2: [sum(stars2.values())]}, index=["Stars"]))

            st.subheader("🧪 Languages Known")
            lang_df = pd.DataFrame({
                user1: pd.Series(langs1),
                user2: pd.Series(langs2)
            }).fillna(0).astype(int)
            st.dataframe(lang_df)

            # Download Reports
            combined_df = pd.DataFrame({
                "Metric": ["Followers", "Following", "Public Repos", "Total Stars", "AI Score"],
                user1: [profile1["followers"], profile1["following"], profile1["public_repos"], 
                       sum(stars1.values()), skill1['score']],
                user2: [profile2["followers"], profile2["following"], profile2["public_repos"], 
                       sum(stars2.values()), skill2['score']],
            })

            st.subheader("📥 Download Comparison Report")
            st.download_button("Download CSV", combined_df.to_csv(index=False), file_name="compare_users.csv")

            # Create PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            pdf.cell(200, 10, txt="GitHub User Comparison Report", ln=True, align='C')
            pdf.ln(10)
            for idx, row in combined_df.iterrows():
                pdf.cell(200, 10, txt=f"{row['Metric']}: {user1} - {row[user1]}, {user2} - {row[user2]}", ln=True)
            pdf_path = "compare_users.pdf"
            pdf.output(pdf_path)
            with open(pdf_path, "rb") as f:
                st.download_button("Download PDF", f, file_name=pdf_path)

            # DOCX
            doc = Document()
            doc.add_heading("GitHub Comparison Report", 0)
            table = doc.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Metric'
            hdr_cells[1].text = user1
            hdr_cells[2].text = user2
            for idx, row in combined_df.iterrows():
                cells = table.add_row().cells
                cells[0].text = str(row["Metric"])
                cells[1].text = str(row[user1])
                cells[2].text = str(row[user2])
            doc_path = "compare_users.docx"
            doc.save(doc_path)
            with open(doc_path, "rb") as f:
                st.download_button("Download DOCX", f, file_name=doc_path)

# ---------- Trending Repositories ----------
elif action == "Trending Repositories":
    st.subheader("🔥 Trending GitHub Repositories")
    trending = [
        {"name": "openai/ChatGPT", "url": "https://github.com/openai/chatgpt"},
        {"name": "facebook/react", "url": "https://github.com/facebook/react"},
        {"name": "microsoft/TypeScript", "url": "https://github.com/microsoft/TypeScript"},
    ]
    for repo in trending:
        st.markdown(f"- [{repo['name']}]({repo['url']})")

# ---------- Trending Developers ----------
elif action == "Trending Developers":
    st.subheader("🌟 Trending GitHub Developers")
    developers = ["torvalds", "gaearon", "JakeWharton"]
    for dev in developers:
        st.markdown(f"- [{dev}](https://github.com/{dev})")

# ---------- Resume Generator ----------
elif action == "Resume Generator":
    resume_user = st.text_input("Enter GitHub Username to Generate Resume")
    if resume_user:
        profile = fetch_profile(resume_user)
        repos, lang_counter, stars, _ = fetch_repos(resume_user)
        if profile:
            # Add AI insights to resume
            st.subheader("🤖 AI-Enhanced Resume Preview")
            skill_data = ai_skill_prediction(lang_counter, repos, profile)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**Expertise Level:**", skill_data['expertise_level'])
                st.write("**AI Score:**", f"{skill_data['score']}/100")
            with col2:
                st.write("**Specializations:**")
                for spec in skill_data['specializations']:
                    st.write(f"- {spec}")
            
            path_docx, path_pdf = generate_resume(profile, repos)
            with open(path_docx, "rb") as f:
                st.download_button("📄 Download Resume (DOCX)", f, file_name=path_docx)
            with open(path_pdf, "rb") as f:
                st.download_button("📄 Download Resume (PDF)", f, file_name=path_pdf)

# ---------- Footer ----------
st.markdown("---")
st.markdown("**🚀 Enhanced with AI/ML Features** | Built with Streamlit & scikit-learn")
                    